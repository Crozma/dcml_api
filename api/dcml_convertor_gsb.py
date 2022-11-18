from docx import Document
from docx.shared import Pt, RGBColor, Cm
from datetime import datetime
from json import loads


class GsbConversionDoc:
    def __init__(self):
        self.meta_data = None
        self.doc = Document()
        self.steckbrief_doc = Document()
        self.counter = 0
        self.line_index = 0
        self.task_number = 0
        self.lektion = None
        self.page_number = None
        self.start_page = None
        self.end_page = None
        self.page_first_call = True
        self.doc_tags = list()
        self.doc_tags_indexes = list()
        self.lines = list()
        self.row = 0
        self.current_par = None
        self.tag_functions = {'page': self.page, 'p': self.neo_text, 'zp': self.line_text,
                              'rm': self.rahmen, 'h1': self.heading_1, 'h2': self.heading_2,
                              'h3': self.heading_3, 'task': self.task, 'list': self.text_list,
                              'fz': self.fusszeile, 'h4': self.heading_4, 'bn': self.bn,
                              'an': self.anmerkung, 'img': self.image}

    @staticmethod
    def word_seperator(body: list, linebreak: bool = True, braces: bool = True) -> list:
        words: list[str] = list()
        final_line = str()

        space_words = [' ', '.', ',', '(', ')']
    
        if len(body) == 1:
            for index, char in enumerate(body[0]):
                if index + 1 == len(body[0]):
                    if char in space_words:
                        if final_line != '':
                            words.append(final_line)
                        words.append(char)
                    else:
                        final_line += char
                        words.append(final_line)
                elif char not in space_words:
                    if braces:
                        if char == '{':
                            final_line += '<'
                        elif char == '}':
                            final_line += '>'
                        else:
                            final_line += char
                    else:
                        final_line += char
                else:
                    if final_line != '':
                        words.append(final_line)
                    final_line = str()
                    if char not in [' ', '']:
                        words.append(char)
        else:
            for line_index, line in enumerate(body):
                for index, char in enumerate(line):
                    if index + 1 == len(line):
                        if char == '{':
                            final_line += '<'
                        elif char == '}':
                            final_line += '>'
                        else:
                            final_line += char
                        if final_line != '':
                            words.append(final_line)
                        final_line = str()
                    elif char not in space_words:
                        if braces:
                            if char == '{':
                                final_line += '<'
                            elif char == '}':
                                final_line += '>'
                            else:
                                final_line += char
                        else:
                            final_line += char
                    else:
                        if final_line != '':
                            words.append(final_line)
                        final_line = str()
                        if char not in [' ', '']:
                            words.append(char)

                if linebreak and line_index != len(body)-1:
                    words.append('\n')

        return words

    def parse(self, file_content):
        word = str()
        words = list()
        tags = list()
        tags_index = list()
        umlaut = False

        with open("api/meta/meta.json", "r") as json_file:
            self.meta_data = loads(json_file.read())

        for char in file_content:
            if char == '\n':
                if word != '':
                    words.append(word.strip())
                    word = ''
            elif char == '>':
                if word != '':
                    word += char
                    words.append(word)
                    word = ''
            elif char == '<':
                if word.strip() != '':
                    words.append(word)
                word = char
            elif char == '*':
                umlaut = True
            elif umlaut:
                if char == 'A':
                    word += 'Ä'
                elif char == 'O':
                    word += 'Ö'
                elif char == 'U':
                    word += 'Ü'
                elif char == 'a':
                    word += 'ä'
                elif char == 'o':
                    word += 'ö'
                elif char == 'u':
                    word += 'ü'
                elif char  == 's':
                    word += 'ß'
                elif char == "'":
                    word += '´'
                umlaut = False
            else:
                word += char

        for index, line in enumerate(words):
            line_list = list(line)

            if line[0] == '<' and line[len(line_list) - 1] == '>':
                if line[1] == '/':
                    content = str(line[1:len(line_list) - 1])
                    tags.append(content)
                    tags_index.append(index)
                else:
                    content = str(line[1:len(line_list) - 1])
                    tags.append(content)
                    tags_index.append(index)

        self.doc_tags = tags

        self.doc_tags_indexes = tags_index
        self.lines = words

    def convert(self):
        for index_start, tag_start in enumerate(self.doc_tags):
            tag_check = True
            tag = str()
            tag_args = list()
            arg = str()

            for char_index, char in enumerate(tag_start):
                if char_index == 0:
                    if char == '/': break

                if char == ' ' or char == '=':
                    if tag_check:
                        tag_check = False
                    else:
                        tag_args.append(arg)
                        arg = str()

                elif char_index >= 1:
                    if char_index == len(tag_start) - 1 and not tag_check:
                        arg += char
                        tag_args.append(arg)
                    else:
                        if tag_check:
                            tag += char
                        else:
                            arg += char
                else:
                    if tag_check:
                        tag += char
                    else:
                        arg += char

            for index_end, tag_end in enumerate(self.doc_tags):
                if index_start < index_end and tag_end == f"/{tag}":
                    tag_body = self.lines[self.doc_tags_indexes[index_start] + 1: self.doc_tags_indexes[index_end]]

                    if not tag_args:
                        self.tag_functions.get(tag)(tag_body)
                    elif tag == 'bn':
                        self.tag_functions.get(tag)()
                    else:
                        self.tag_functions.get(tag)(tag_body, args=tag_args)

                    break

    def page(self, body: list, args: list = None):
        doc = self.doc
        page_number = 0

        if args is not None:
            if self.page_first_call:
                page_number = args[1]

            if self.end_page is None:
                self.end_page = int(page_number)

        else:
            page_number = str(int(self.page_number) + 1)
            self.end_page += 1

        doc.add_paragraph()
        doc.add_paragraph()
        number = doc.add_paragraph().add_run(f'(({page_number}))')
        font = number.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = self.meta_data.get("text").get("groeße")

        if self.page_first_call is True:
            self.start_page = page_number

        self.page_number = page_number
        self.page_first_call = False

    def text(self, body: list, paragraph=None, linebreak=False):
        doc = self.doc
        text = str()
        final_line = str()

        if len(body) == 1:
            for char in body[0]:
                if char == '{':
                    final_line += '<'
                elif char == '}':
                    final_line += '>'
                else:
                    final_line += char
            text = final_line

        else:
            for index, line in enumerate(body):
                final_line = str()
                for char in line:
                    if char == '{':
                        final_line += '<'
                    elif char == '}':
                        final_line += '>'
                    else:
                        final_line += char
                if linebreak and index != 0:
                    if index != len(body) - 1:
                        text += final_line + ' '
                    else:
                        text = final_line
                else:
                    if index != len(body) - 1:
                        text += final_line + '\n'
                    else:
                        text += final_line

        if paragraph is None:
            if self.current_par is None:
                run = doc.add_paragraph().add_run(text)
                font = run.font
                font.name = self.meta_data.get("text").get("schrift")
                font.size = Pt(self.meta_data.get("text").get("size"))
            else:
                run = self.current_par.add_run('\n' + text)
                font = run.font
                font.name = self.meta_data.get("text").get("schrift")
                font.size = Pt(self.meta_data.get("text").get("size"))

                self.current_par = None
        else:
            run = paragraph.add_run(text)
            font = run.font
            font.name = self.meta_data.get("text").get("schrift")
            font.size = Pt(self.meta_data.get("text").get("size"))

        return run

    def neo_text(self, body: list, args: list = None, paragraph=None, linebreak=True, braces=True, add_linebreak=False):
        doc = self.doc
        non_space_words = ('\n', '.', ',', ':', ')')

        if args is not None:
            if 'lb' in args:
                linebreak = False
            if 'bc' in args:
                braces = False

        words = self.word_seperator(body, linebreak, braces)

        if paragraph is None:
            if self.current_par is None:
                par = doc.add_paragraph()
            else:
                par = self.current_par
                self.current_par.add_run('\n')
        else:
            par = paragraph

        bold_check = False
        italic_check = False
        underline_check = False

        for index, word in enumerate(words):
            if word == '#b':
                bold_check = True
            elif word == '#i':
                italic_check = True
            elif word == '#u':
                underline_check = True
            else:
                run = par.add_run(word)
                font = run.font
                font.name = self.meta_data.get("text").get("schrift")
                font.size = Pt(self.meta_data.get("text").get("size"))
                if bold_check:
                    font.bold = True
                    bold_check = False
                if italic_check:
                    font.italic = True
                    italic_check = False
                if underline_check:
                    run.underline = True
                    underline_check = False

                if index + 1 < len(words):
                    if words[index+1] not in non_space_words and word not in ['(', '\n', '']:
                        par.add_run(' ')

        if paragraph is not None and add_linebreak:
            paragraph.add_run('\n')
        self.current_par = None

    def rahmen(self, body: list, ):
        doc = self.doc
        self.current_par = None

        par = doc.add_paragraph()

        self.text(['<Rahmen>\n'], paragraph=par)
        self.neo_text([body[0]], paragraph=par, add_linebreak=True)
        self.neo_text(body[1:], paragraph=par, linebreak=False)
        self.text(['\n</Rahmen>'], paragraph=par)

    def heading_1(self, body: list):
        doc = self.doc
        self.current_par = doc.add_paragraph()

        text = self.text(body, paragraph=self.current_par)

        font = text.font
        font.name = self.meta_data.get("heading 1").get("schrift")
        font.size = Pt(self.meta_data.get("heading 1").get("size"))
        font.color.rgb = RGBColor(54, 96, 145)
        font.bold = True

    def heading_2(self, body: list):
        doc = self.doc
        self.current_par = doc.add_paragraph()

        text = self.text(body, paragraph=self.current_par)

        font = text.font
        font.name = self.meta_data.get("heading 2").get("schrift")
        font.size = Pt(self.meta_data.get("heading 2").get("size"))
        font.color.rgb = RGBColor(79, 129, 189)
        font.bold = True

    def heading_3(self, body: list):
        doc = self.doc
        self.current_par = doc.add_paragraph()

        text = self.text(body, paragraph=self.current_par)

        font = text.font
        font.name = self.meta_data.get("heading 3").get("schrift")
        font.size = Pt(self.meta_data.get("heading 3").get("size"))
        font.color.rgb = RGBColor(79, 129, 189)
        font.bold = True

    def heading_4(self, body: list):
        doc = self.doc
        self.current_par = doc.add_paragraph()

        text = self.text(body, paragraph=self.current_par)

        font = text.font
        font.name = self.meta_data.get("heading 4").get("schrift")
        font.size = Pt(self.meta_data.get("heading 4").get("size"))
        font.color.rgb = RGBColor(79, 129, 189)
        font.bold = True

    def task(self, body: list, args: list = None):
        task_number = self.task_number

        if args is not None:
            task_number = int(args[args.index('num') + 1])
            if len(body) == 1:
                self.text([f'{task_number}. {body[0]}'], linebreak=True)
            else:
                self.text([f'{task_number}. {body[0]}'] + body[1:], linebreak=True)

        else:
            task_number += 1
            if len(body) == 1:
                self.text([f'{task_number}. {body[0]}'], linebreak=True)
            else:
                self.text([f'{task_number}. {body[0]}'] + body[1:], linebreak=True)

        self.task_number = task_number
        self.current_par = None

    def text_list(self, body: list[str]):
        doc = self.doc

        for line in body:
            final_line = str()
            if line.startswith('- '):
                par = doc.add_paragraph(style='List Bullet 3')
                final_line = line[2:]
            else:
                final_line = line
                par = doc.add_paragraph(style="List Bullet 2")
            self.text([final_line], paragraph=par)

        self.current_par = None

    def line_text(self, body: list, args: list = None, line_mark: int = 5):
        doc = self.doc
        counter = self.counter

        if args is not None:
            if args[args.index('type') + 1] == 'start':
                line_index = 0
                counter = 0
            else:
                line_index = self.line_index
        else:
            line_index = self.line_index

        for line in body:
            line_index += 1
            counter += 1
            check = True

            par = doc.add_paragraph()

            for index, char in enumerate(line):
                if counter == line_mark:
                    if check:
                        run = par.add_run(f'{line_index}\t{char}')
                        check = False
                    else:
                        run = par.add_run(char)
                    font = run.font
                    font.name = self.meta_data.get("text").get("schrift")
                    font.size = Pt(self.meta_data.get("text").get("size"))
                else:
                    run = par.add_run(char)
                    par.paragraph_format.left_indent = Cm(0.5)
                    font = run.font
                    font.name = self.meta_data.get("text").get("schrift")
                    font.size = Pt(self.meta_data.get("text").get("size"))
            if counter == 5:
                counter = 0

        self.line_index = line_index
        self.counter = counter
        self.current_par = None

    def fusszeile(self, body: list):
        doc = self.doc

        self.text(['Fußzeile:'])

        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.5)

        for index, line in enumerate(body):
            if index == 0:
                self.text([f'^{index + 1} {line}'], paragraph=par)
            else:
                self.text([f'\n^{index + 1} {line}'], paragraph=par)
        self.current_par = None

    def bn(self, body: list):
        self.text(['<Anmerkung>', 'Bild nicht übertragen', '</Anmerkung>'])
        self.current_par = None

    def anmerkung(self, body: list):
        doc = self.doc

        par = doc.add_paragraph()

        self.text(['{Anmerkung}\n'], paragraph=par)
        self.neo_text(body, paragraph=par, linebreak=False)
        self.text(['\n{/Anmerkung}'], paragraph=par)

    def image(self, body: list):
        doc = self.doc

        par = doc.add_paragraph()

        
        self.text(['{Bild}\n'], paragraph=par)
        self.neo_text(body, paragraph=par, linebreak=False)
        self.text(['\n{/Bild}'], paragraph=par)
 
    def create_steckbrief(self):
        steckbrief = self.steckbrief_doc

        par = steckbrief.add_paragraph()

        heading = par.add_run('Diese Steckbriefdatei\n')
        font = heading.font
        font.name = 'Calibri Bold'
        font.size = Pt(14)
        font.color.rgb = RGBColor(54, 96, 145)
        font.bold = True

        run = par.add_run(f'E-Buch Steckbrief, P.A.U.L. D. 10 {self.start_page}-{self.end_page}')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        heading = par.add_run('Quelle')
        font = heading.font
        font.name = 'Calibri Bold'
        font.size = Pt(14)
        font.color.rgb = RGBColor(54, 96, 145)
        font.bold = True

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run('Titel: P.A.U.L. D. 10')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run('Autor: Martin Fruhstorfer')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run('Jahr: 2019')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run('Verlag: Schöningh')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run('ISBN: 978-3-14-127420-2')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        heading = par.add_run('Übertragung')
        font = heading.font
        font.name = 'Calibri Bold'
        font.size = Pt(14)
        font.color.rgb = RGBColor(54, 96, 145)
        font.bold = True

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run(f'Übertragene Seiten: {self.start_page}-'
                                                                      f'{self.end_page}')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run('Übertragen von: Michael Lukas')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        run = steckbrief.add_paragraph(style='List Bullet 2').add_run(f'{str(datetime.now().day)}.'
                                                                      f'{str(datetime.now().month)}.'
                                                                      f'{str(datetime.now().year)}')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        heading = par.add_run('Verwendete Formatvorlage\n')
        font = heading.font
        font.name = 'Calibri Bold'
        font.size = Pt(14)
        font.color.rgb = RGBColor(54, 96, 145)
        font.bold = True

        run = par.add_run(f'Überschrift 1\n')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)
        font.bold = True

        run = par.add_run(f'\tInhaltsverzeichnis (Überschrift)')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        run = par.add_run(f'Überschrift 2\n')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)
        font.bold = True

        run = par.add_run(f'\tHauptkapitelüberschrift, z.B. "So fangen Romane an"')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        run = par.add_run(f'Überschrift 3\n')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)
        font.bold = True

        run = par.add_run(f'\tUnterkapitelüberschrift, z.B. "1. Erzählform, Erzählerstandort und \tErzählverhalten 20"')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        run = par.add_run(f'Überschrift 4\n')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)
        font.bold = True

        run = par.add_run(f'\tTexte innerhalb des Kapitels z.B. "Erich Maria Remarque: Im Westen '
                          f'\tnichts Neues (Romanauszug) 20"')

        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

        par = steckbrief.add_paragraph()

        heading = par.add_run('Verwendete E-Buch-Tags\n')
        font = heading.font
        font.name = 'Calibri Bold'
        font.size = Pt(14)
        font.color.rgb = RGBColor(54, 96, 145)
        font.bold = True

        run = par.add_run(f'<Anmerkung> ... </Anmerkung>\n')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)
        font.bold = True

        run = par.add_run(f'\tAnmerkungen zur Übertragung wie z.B. Beschreibung von Bildern,')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)
        run = par.add_run(f'\tTabellen, usw.')
        font = run.font
        font.name = self.meta_data.get("text").get("schrift")
        font.size = Pt(11)

    def save(self, directory: str):
        self.doc.save(directory + f'((0{self.start_page}-0{self.end_page})) Paul D 10 Schülerbuch '
                                  f'978-3-14-127420-2.docx')
        self.create_steckbrief()
        self.steckbrief_doc.save(directory + f'((0{self.start_page}-0{self.end_page})) E-Buch Steckbrief, '
                                             f'P.A.U.L. D. 10.docx')

        print('Saved Documents successfully')
